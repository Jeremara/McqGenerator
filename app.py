import streamlit as st
import json
import os
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import TextLoader, PyPDFLoader
from langchain.vectorstores import FAISS
from langchain_openai import AzureChatOpenAI
from langchain.embeddings import AzureOpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import tempfile
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="RAG MCQ Question Generator",
    page_icon="📝",
    layout="wide"
)

# Initialize session state
if 'questions_generated' not in st.session_state:
    st.session_state.questions_generated = False
if 'mcq_questions' not in st.session_state:
    st.session_state.mcq_questions = []
if 'short_questions' not in st.session_state:
    st.session_state.short_questions = []
if 'show_answers' not in st.session_state:
    st.session_state.show_answers = {}

def initialize_models():
    try:
        embedding_model = AzureOpenAIEmbeddings(
            azure_endpoint=os.getenv("EMBEDDING_API_BASE"),
            openai_api_key=os.getenv("EMBEDDING_API_KEY"),
            deployment=os.getenv("EMBEDDING_DEPLOYMENT_NAME"),
            chunk_size=10,
        )

        llm = AzureChatOpenAI(
            azure_endpoint=os.getenv("AZURE_OPENAI_API_BASE"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
            azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
            temperature=0.7,
            max_tokens=800,
            model_name="gpt-4o"
        )

        return embedding_model, llm
    except Exception as e:
        st.error(f"Error initializing models: {str(e)}")
        return None, None

def process_document(file, file_type):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_file_path = tmp_file.name

        loader = PyPDFLoader(tmp_file_path) if file_type == "pdf" else TextLoader(tmp_file_path)
        documents = loader.load()
        os.unlink(tmp_file_path)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        return text_splitter.split_documents(documents)

    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        return None

def generate_questions(split_docs, embedding_model, llm, question_type):
    try:
        vector_db = FAISS.from_documents(split_docs, embedding_model)
        retriever = vector_db.as_retriever()

        prompt = """
        Generate 5 MCQ questions in contextx with 4 options each (A, B, C, D) and indicate the correct answer.
        Return the response in valid JSON format like this:
        {
        "questions": [
            {
            "question": "What is...?",
            "options": {
                "A": "Option A",
                "B": "Option B",
                "C": "Option C",
                "D": "Option D"
            },
            "correct_answer": "A",
            "explanation": "Explanation here"
            }
        ]
        }
        """ if question_type == "mcq" else """
        Generate 5 Short answer questions in context with detailed answers.
        Return the response in valid JSON format like this:
        {
        "questions": [
            {
            "question": "What is...?",
            "answer": "Detailed answer here",
            "key_points": ["Point 1", "Point 2"]
            }
        ]
        }
        """

        qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
        response = qa_chain.invoke(prompt)
        return response["result"]

    except Exception as e:
        st.error(f"Error generating questions: {str(e)}")
        return None

def parse_questions_response(response):
    try:
        response = response.strip()
        start_idx = response.find('{')
        end_idx = response.rfind('}') + 1
        if start_idx != -1 and end_idx != 0:
            json_str = response[start_idx:end_idx]
            questions_data = json.loads(json_str)
            return questions_data.get("questions", [])
        else:
            st.error("No valid JSON found in response")
            return []
    except json.JSONDecodeError as e:
        st.error(f"JSON parsing error: {str(e)}")
        st.error(f"Response received: {response[:500]}...")
        return []
    except Exception as e:
        st.error(f"Error processing response: {str(e)}")
        return []

def display_mcq_questions(questions):
    for i, q in enumerate(questions):
        st.markdown(f"**Question {i+1}:** {q.get('question', '')}")
        options = q.get('options', {})
        for key, val in options.items():
            st.markdown(f"- {key}) {val}")

        answer_key = f"mcq_answer_{i}"
        if st.button(f"Show Answer", key=f"show_mcq_{i}"):
            st.session_state.show_answers[answer_key] = not st.session_state.show_answers.get(answer_key, False)

        if st.session_state.show_answers.get(answer_key, False):
            st.markdown(f"**Correct Answer:** {q.get('correct_answer', '')}")
            st.markdown(f"**Explanation:** {q.get('explanation', '')}")
        st.markdown("---")

def display_short_questions(questions):
    for i, q in enumerate(questions):
        st.markdown(f"**Question {i+1}:** {q.get('question', '')}")
        answer_key = f"short_answer_{i}"
        if st.button(f"Show Answer", key=f"show_short_{i}"):
            st.session_state.show_answers[answer_key] = not st.session_state.show_answers.get(answer_key, False)

        if st.session_state.show_answers.get(answer_key, False):
            st.markdown(f"**Answer:** {q.get('answer', '')}")
            key_points = q.get("key_points", [])
            if key_points:
                st.markdown("*Key Points:*")
                for point in key_points:
                    st.markdown(f"- {point}")
        st.markdown("---")

def generate_docx(mcqs, shorts):
    doc = Document()
    doc.add_heading("Generated Questions", 0)

    if mcqs:
        doc.add_heading("Multiple Choice Questions", level=1)
        for i, q in enumerate(mcqs, 1):
            doc.add_paragraph(f"{i}. {q.get('question', '')}")
            for key, val in q.get("options", {}).items():
                doc.add_paragraph(f"{key}) {val}", style="List Bullet")
            doc.add_paragraph(f"Correct Answer: {q.get('correct_answer', '')}")
            doc.add_paragraph(f"Explanation: {q.get('explanation', '')}")

    if shorts:
        doc.add_heading("Short Answer Questions", level=1)
        for i, q in enumerate(shorts, 1):
            doc.add_paragraph(f"{i}. {q.get('question', '')}")
            doc.add_paragraph(f"Answer: {q.get('answer', '')}")
            if q.get("key_points"):
                doc.add_paragraph("Key Points:")
                for pt in q["key_points"]:
                    doc.add_paragraph(f"• {pt}", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(mcqs, shorts):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Generated Questions", ln=True)
    pdf.set_font("Arial", "", 12)

    if mcqs:
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Multiple Choice Questions", ln=True)
        pdf.set_font("Arial", "", 12)
        for i, q in enumerate(mcqs, 1):
            pdf.multi_cell(0, 10, f"{i}. {q.get('question', '')}")
            for key, val in q.get("options", {}).items():
                pdf.multi_cell(0, 10, f"  {key}) {val}")
            pdf.multi_cell(0, 10, f"Correct Answer: {q.get('correct_answer', '')}")
            pdf.multi_cell(0, 10, f"Explanation: {q.get('explanation', '')}")
            pdf.ln()

    if shorts:
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Short Answer Questions", ln=True)
        pdf.set_font("Arial", "", 12)
        for i, q in enumerate(shorts, 1):
            pdf.multi_cell(0, 10, f"{i}. {q.get('question', '')}")
            pdf.multi_cell(0, 10, f"Answer: {q.get('answer', '')}")
            if q.get("key_points"):
                pdf.multi_cell(0, 10, "Key Points:")
                for pt in q["key_points"]:
                    pdf.multi_cell(0, 10, f"  - {pt}")
            pdf.ln()

    # Return PDF as BytesIO buffer
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    return BytesIO(pdf_bytes)

def show_download_button():
    mcqs = st.session_state.mcq_questions
    shorts = st.session_state.short_questions
    if mcqs or shorts:
        docx_file = generate_docx(mcqs, shorts)
        pdf_file = generate_pdf(mcqs, shorts)

        st.download_button(
            label="📄 Download as DOCX",
            data=docx_file,
            file_name="generated_questions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.download_button(
            label="📄 Download as PDF",
            data=pdf_file,
            file_name="generated_questions.pdf",
            mime="application/pdf"
        )

def main():
    st.title("📝 RAG MCQ Question Generator")
    st.markdown("Upload a document and generate MCQ or Short Answer questions")

    with st.sidebar:
        st.header("Upload Document")
        uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'txt'])
        if uploaded_file:
            file_type = uploaded_file.name.split('.')[-1].lower()
            st.success(f"Uploaded: {uploaded_file.name}")

            if st.button("Generate Questions"):
                with st.spinner("Processing..."):
                    embedding_model, llm = initialize_models()
                    if embedding_model and llm:
                        split_docs = process_document(uploaded_file, file_type)
                        if split_docs:
                            mcq_response = generate_questions(split_docs, embedding_model, llm, "mcq")
                            short_response = generate_questions(split_docs, embedding_model, llm, "short")
                            st.session_state.mcq_questions = parse_questions_response(mcq_response)
                            st.session_state.short_questions = parse_questions_response(short_response)
                            st.session_state.questions_generated = True
                            st.success("Questions generated!")
                            st.rerun()

    if st.session_state.questions_generated:
        tab1, tab2 = st.tabs(["Multiple Choice Questions", "Short Answer Questions"])
        with tab1:
            st.subheader("Multiple Choice Questions")
            display_mcq_questions(st.session_state.mcq_questions)
        with tab2:
            st.subheader("Short Answer Questions")
            display_short_questions(st.session_state.short_questions)

        show_download_button()
        if st.button("🔄 Reset and Upload New Document"):
            st.session_state.questions_generated = False
            st.session_state.mcq_questions = []
            st.session_state.short_questions = []
            st.session_state.show_answers = {}
            st.rerun()
    else:
        st.info("👈 Upload a document from the sidebar to get started.")

if __name__ == "__main__":
    main()
